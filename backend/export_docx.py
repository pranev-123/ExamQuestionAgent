from io import BytesIO

from docx import Document


def build_docx(content: str, title: str = "Generated Question Paper") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.splitlines():
        if line.strip():
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
